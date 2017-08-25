import docx
import datetime
import getArgs
import os
import time


global pro_name
global variables
variables = getArgs.getArgs("项目相关变量.txt")
pro_name = variables["项目名称"]
pro_type = variables["项目类型"]

def _log(string):
    def outter(fun):
        with open(r"C:\Users\Administrator\Documents\WordLog.txt",'a') as f:
            f.write(str(datetime.datetime.now())+"  -> "+string+"\n")
        def wrapper(self):
            fun(self)
            print("End")
        return wrapper
    return outter


class WordSub:
    def __init__(self,*,word_file_name,out_file_name):
        self.word_file = docx.Document(word_file_name)
        self.out_file_name = out_file_name
        self.count = len(self.word_file.paragraphs)


    def __substitute(self,var):
        if var.strip("{}") not in variables.keys():
            raise Exception("请查看一下是否变量名称错误->"+var)
        return variables[var.strip("{}")]

    @_log(pro_name)
    def run(self):
        print("wait for ",str(self.out_file_name))
        for num,para in enumerate(self.word_file.paragraphs):
            print("Processing:",num,'/',self.count,end="  ")
            global var
            global inStr
            var = ""
            inStr = False
            for r in para.runs:
                if r.text.startswith("{") and not r.text.endswith("}"):
                    inStr=True
                    var+=r.text
                    r.text=""
                elif inStr and not r.text.endswith("}"):
                    var+=r.text
                    r.text=""
                elif inStr and r.text.endswith("}"):
                    var+=r.text
                    r.text=self.__substitute(var)
                    inStr = False
                    var = ""
                elif r.text.startswith("{") and r.text.endswith("}"):
                    r.text = self.__substitute(r.text)
            print()

        # Update the tables
        for table in self.word_file.tables:
            var = ""
            inStr = False
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for r in para.runs:
                            if r.text.startswith("{") and not r.text.endswith("}"):
                                inStr = True
                                var += r.text
                                r.text = ""
                            elif inStr and not r.text.endswith("}"):
                                var += r.text
                                r.text = ""
                            elif inStr and r.text.endswith("}"):
                                var += r.text
                                r.text = self.__substitute(var)
                                inStr = False
                                var = ""
                            elif r.text.startswith("{") and r.text.endswith("}"):
                                r.text = self.__substitute(r.text)
        self.word_file.save(self.out_file_name)




if __name__ == "__main__":
    try:
        beg = time.time()
        wb = WordSub(word_file_name="./template/{项目名称}{项目类型}采购合同模板.docx", out_file_name = pro_name + pro_type+ "采购合同.docx")
        ws = WordSub(word_file_name="./template/{项目名称}{项目类型}销售合同模板.docx", out_file_name = pro_name + pro_type+ "销售合同.docx")
        wc = WordSub(word_file_name="./template/{项目名称}{项目类型}协作合同模板.docx", out_file_name = pro_name + pro_type+ "协作合同.docx")
        print("开始完成销售合同")
        ws.run()
        print("开始完成采购合同")
        wb.run()
        print("开始完成协作合同")
        wc.run()
        end = time.time()
        print("共花费了"+str(end-beg)+"秒钟的时间")
    except Exception as err:
        print("\n\n有错误发生了，详见下列说明，若有问题请联系孙博\n",str(err))
    os.system("pause")
    
